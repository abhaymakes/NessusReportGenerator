# from selenium import webdriver
# from selenium.webdriver.common.by import By
# from selenium.webdriver.chrome.options import Options
# from selenium.webdriver.support.wait import WebDriverWait
# from selenium.webdriver.support import expected_conditions as EC
# import time
# from bs4 import BeautifulSoup
# from PIL import Image, ImageDraw, ImageFont
# import os

# from docx import Document
# from docx.shared import Pt, RGBColor, Inches, Cm
# from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
# from docx.oxml import OxmlElement
# from docx.oxml.ns import qn

# # --- Selenium setup ---
# options = Options()
# options.add_argument("user-data-dir=C:\\Abhay\\NessusGeneratorProfile")  # Your custom profile path here

# URL = 'https://localhost:8834/#/scans/reports/8/vulnerabilities'  # Change to your scan report URL

# # --- Function to create code block image ---
# def text_to_image(
#     text,
#     font_path="dejavu-sans-mono.book.ttf",  # Change if you want a different font
#     font_size=18,
#     padding=14,
#     bg_color="#f0f0f0",
#     text_color="#111111",
#     scale=2,
#     image_width=1080,
# ):
#     font_size *= scale
#     padding *= scale
#     image_width *= scale

#     font = ImageFont.truetype(font_path, font_size)
#     lines = text.splitlines() or [""]

#     dummy = Image.new("RGBA", (1, 1))
#     draw = ImageDraw.Draw(dummy)
#     line_heights = []

#     for line in lines:
#         bbox = draw.textbbox((0, 0), line, font=font)
#         height = bbox[3] - bbox[1]
#         line_heights.append(height)

#     line_height = max(line_heights) + int(font_size * 0.2)
#     img_height = line_height * len(lines) + padding * 2
#     img_width = image_width 

#     img = Image.new("RGBA", (img_width, img_height), bg_color)
#     draw = ImageDraw.Draw(img)

#     y = padding
#     for line in lines:
#         draw.text((padding, y), line, font=font, fill=text_color)
#         y += line_height

#     img = img.resize(
#         (img.width // scale, img.height // scale),
#         Image.Resampling.LANCZOS
#     )

#     return img.convert("RGB")

# # --- Word formatting helper functions ---
# def add_horizontal_line(paragraph):
#     p = paragraph._p
#     pPr = p.get_or_add_pPr()
#     pBdr = OxmlElement('w:pBdr')
#     bottom = OxmlElement('w:bottom')
#     bottom.set(qn('w:val'), 'single')
#     bottom.set(qn('w:sz'), '4')
#     bottom.set(qn('w:space'), '1')
#     bottom.set(qn('w:color'), 'A6A6A6')
#     pBdr.append(bottom)
#     pPr.append(pBdr)

# def add_section_heading(doc, text):
#     p = doc.add_paragraph(text)
#     p.style.font.name = 'Garamond'
#     p.style.font.size = Pt(11)
#     p.style.font.bold = True
#     p.style.font.color.rgb = RGBColor(0, 0, 0)
#     add_horizontal_line(p)
#     return p

# def add_normal_text(doc, text, space_after=True):
#     p = doc.add_paragraph(text)
#     p.style.font.name = 'Garamond'
#     p.style.font.size = Pt(10)
#     p.style.font.color.rgb = RGBColor(0, 0, 0)
#     if space_after:
#         p.space_after = Pt(8)
#     return p

# def add_vulnerability_to_word_nessus_style(doc, data, image_path=None):
#     severity_colors = {
#         "Critical": ("7B2222", "FFFFFF"),  # Dark red bg, white fg
#         "High": ("9C3B00", "FFFFFF"),      # Dark orange bg, white fg
#         "Medium": ("9C6800", "000000"),    # Dark yellow bg, black fg
#         "Low": ("6C6C6C", "FFFFFF"),       # Grey bg, white fg
#         "Info": ("2D6DB2", "FFFFFF")       # Blue bg, white fg
#     }

#     severity = data.get("severity", "Info")
#     bg_color, fg_color = severity_colors.get(severity, ("2D6DB2", "FFFFFF"))

#     # Vulnerability title banner
#     banner = doc.add_paragraph()
#     run = banner.add_run(f"{data.get('plugin_id', '')} - {data.get('title', '')}")
#     font = run.font
#     font.name = 'Garamond'
#     font.size = Pt(12)
#     font.bold = True
#     font.color.rgb = RGBColor.from_string(fg_color)
#     banner.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

#     shading_elm = OxmlElement('w:shd')
#     shading_elm.set(qn('w:fill'), bg_color)
#     banner._p.get_or_add_pPr().append(shading_elm)

#     doc.add_paragraph()

#     # Sections to add if exist
#     def add_section(title, content):
#         if not content or content.strip() == "":
#             return
#         add_section_heading(doc, title)
#         add_normal_text(doc, content)

#     # Use 'synopsis' if available, else fallback to description for synopsis section
#     add_section("Synopsis", data.get("synopsis", data.get("description", "")))

#     # Avoid duplicate if same
#     if "description" in data and data["description"] != data.get("synopsis", ""):
#         add_section("Description", data["description"])

#     if "see also" in data:
#         add_section("See Also", data["see also"])

#     if "solution" in data:
#         add_section("Solution", data["solution"])

#     if "risk factor" in data:
#         add_section("Risk Factor", data["risk factor"])

#     if "cvss v3.0 base score" in data:
#         add_section("CVSS v3.0 Base Score", data["cvss v3.0 base score"])

#     if "cvss v2.0 base score" in data:
#         add_section("CVSS v2.0 Base Score", data["cvss v2.0 base score"])

#     if "plugin information" in data:
#         add_section("Plugin Information", data["plugin information"])

#     if "plugin output" in data:
#         add_section("Plugin Output", data["plugin output"])

#     # Add POC/code output image if present
#     if image_path and os.path.exists(image_path):
#         doc.add_paragraph()
#         p_img_title = doc.add_paragraph("Proof of Concept / Code Output")
#         p_img_title.style.font.name = 'Garamond'
#         p_img_title.style.font.size = Pt(11)
#         p_img_title.style.font.bold = True
#         p_img_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
#         doc.add_paragraph()
#         doc.add_picture(image_path, width=Inches(6))
#         doc.add_paragraph()

# # --- MAIN SCRIPT ---

# driver = webdriver.Chrome(options=options)
# wait = WebDriverWait(driver, 10)
# driver.get(URL)

# doc = Document()
# # Set margins
# sections = doc.sections
# for section in sections:
#     section.top_margin = Cm(2)
#     section.bottom_margin = Cm(2)
#     section.left_margin = Cm(2.5)
#     section.right_margin = Cm(2.5)

# # Set default style font to Garamond
# style = doc.styles['Normal']
# font = style.font
# font.name = 'Garamond'
# font.size = Pt(11)

# try:
#     time.sleep(2)

#     if 'Login' in driver.title:
#         username_field = driver.find_element(By.CLASS_NAME, 'login-username')
#         password_field = driver.find_element(By.CLASS_NAME, 'login-password')
#         remember_me = driver.find_element(By.CLASS_NAME, 'login-remember')
#         sign_in_button = driver.find_element(By.TAG_NAME, 'button')

#         username_field.clear()
#         username_field.send_keys('admin')  # Change your username here
#         password_field.send_keys('admin')  # Change your password here
#         remember_me.click()
#         sign_in_button.click()

#     element = wait.until(EC.visibility_of_element_located((By.CLASS_NAME, 'select2-selection__rendered')))

#     show_200_results_span = driver.find_element(By.CLASS_NAME, 'select2-selection__rendered')
#     show_200_results_span.click()
#     time.sleep(1)
#     option_200 = driver.find_elements(By.CLASS_NAME, 'select2-results__option')[-1]
#     option_200.click()
#     time.sleep(1)

#     list_of_vulnerabilities = driver.find_elements(By.CSS_SELECTOR, 'tr.vulnerability.add-plugin-id-tip')
#     urls_to_visit = []

#     for vulnerability in list_of_vulnerabilities:
#         urls_to_visit.append(f"{URL}/{vulnerability.get_attribute('data-id')}")

#     for index, url in enumerate(urls_to_visit):
#         data_dictionary = {}

#         driver.get(url)
#         element = wait.until(EC.visibility_of_element_located((By.CLASS_NAME, 'add-plugin-id-tip')))

#         section_data = driver.find_element(By.ID, 'content')
#         section_html = section_data.get_attribute('innerHTML')

#         soup = BeautifulSoup(section_html, 'html.parser')
#         plugin_id = soup.find('tr', class_='noaction odd').get('data-plugin-id')
#         vuln_severity = soup.find('span', class_='severity').text
#         vuln_title = soup.find('h4', class_='add-plugin-id-tip').text

#         data_dictionary['plugin_id'] = plugin_id
#         data_dictionary['severity'] = vuln_severity
#         data_dictionary['title'] = vuln_title

#         vuln_data = soup.find_all('div', 'plugin-details-content')[0]

#         # Map headers to dictionary keys in lowercase, with spaces replaced by underscores for consistency
#         for header, content in zip(vuln_data.find_all('h5'), vuln_data.find_all('div', 'plugin-wrap')):
#             key = header.text.lower()
#             # normalize key by replacing spaces with underscores for easier access (optional)
#             key = key.replace(" ", "_")
#             data_dictionary[key] = content.text.strip()

#         plugin_output_code = soup.find('pre', class_='monospace').text
#         data_dictionary['plugin_output'] = plugin_output_code

#         port_and_host = soup.find('tr', class_='noaction odd').find_all('td')
#         port, host = port_and_host[0].text.strip(), port_and_host[1].text.strip()

#         data_dictionary['port'] = port
#         data_dictionary['host'] = host

#         # Save image for code output
#         image_filename = f"{plugin_id}.png"
#         codeblock_image = text_to_image(plugin_output_code)
#         codeblock_image.save(image_filename)

#         # Append vulnerability to Word doc
#         add_vulnerability_to_word_nessus_style(doc, data_dictionary, image_path=image_filename)

#         # Clean up image file after use
#         os.remove(image_filename)

#         print(f"Processed vulnerability {index+1}/{len(urls_to_visit)}", end='\r', flush=True)

#     # Save final Word document
#     doc.save("Nessus_Report.docx")
#     print("\nReport saved as Nessus_Report.docx")

# except Exception as e:
#     print(f"An error occurred: {e}")

# finally:
#     driver.quit()


from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.wait import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import time
from bs4 import BeautifulSoup
from PIL import Image, ImageDraw, ImageFont
import os

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement

# --- Selenium setup ---
options = Options()
options.add_argument("user-data-dir=C:\\Abhay\\NessusGeneratorProfile")  # Your profile path

URL = 'https://localhost:8834/#/scans/reports/8/vulnerabilities'  # Your URL

def text_to_image(
    text,
    font_path="dejavu-sans-mono.book.ttf",
    font_size=18,
    padding=14,
    bg_color="#f0f0f0",
    text_color="#111111",
    scale=2,
    image_width=1080,
):
    font_size *= scale
    padding *= scale
    image_width *= scale

    font = ImageFont.truetype(font_path, font_size)
    lines = text.splitlines() or [""]

    dummy = Image.new("RGBA", (1, 1))
    draw = ImageDraw.Draw(dummy)
    line_heights = []

    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        height = bbox[3] - bbox[1]
        line_heights.append(height)

    line_height = max(line_heights) + int(font_size * 0.2)
    img_height = line_height * len(lines) + padding * 2
    img_width = image_width 

    img = Image.new("RGBA", (img_width, img_height), bg_color)
    draw = ImageDraw.Draw(img)

    y = padding
    for line in lines:
        draw.text((padding, y), line, font=font, fill=text_color)
        y += line_height

    img = img.resize(
        (img.width // scale, img.height // scale),
        Image.Resampling.LANCZOS
    )

    return img.convert("RGB")

def add_horizontal_line(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'A6A6A6')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_section_heading(doc, text):
    p = doc.add_paragraph(text)
    p.style.font.name = 'Arial'
    p.style.font.size = Pt(10)
    p.style.font.color.rgb = RGBColor(0, 0, 0)
    # No bold
    add_horizontal_line(p)
    return p

def add_normal_text(doc, text, space_after=True):
    p = doc.add_paragraph(text)
    p.style.font.name = 'Arial'
    p.style.font.size = Pt(10)
    p.style.font.color.rgb = RGBColor(0, 0, 0)
    if space_after:
        p.space_after = Pt(6)
    return p

def add_vulnerability_to_word_nessus_style(doc, data, image_path=None):
    severity_colors = {
        "Critical": ("6D203E", "FFFFFF"),  # Wine red bg, white fg per your image
        "High": ("9C3B00", "FFFFFF"),      # fallback colors if needed
        "Medium": ("9C6800", "000000"),
        "Low": ("6C6C6C", "FFFFFF"),
        "Info": ("2D6DB2", "FFFFFF")
    }

    severity = data.get("severity", "Info")
    bg_color, fg_color = severity_colors.get(severity, ("2D6DB2", "FFFFFF"))

    # Vulnerability title banner - with padding inside the banner
    banner = doc.add_paragraph()
    run = banner.add_run(f"{data.get('plugin_id', '')} - {data.get('title', '')}")
    font = run.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor.from_string(fg_color)
    # no bold

    banner.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    # Add shading with padding effect: use indent & spacing to simulate padding
    banner.paragraph_format.left_indent = Pt(6)
    banner.paragraph_format.right_indent = Pt(6)
    banner.paragraph_format.space_before = Pt(6)
    banner.paragraph_format.space_after = Pt(6)
    banner.paragraph_format.keep_with_next = True

    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), bg_color)
    banner._p.get_or_add_pPr().append(shading_elm)

    # Sections to add if exist
    def add_section(title, content):
        if not content or content.strip() == "":
            return
        add_section_heading(doc, title)
        add_normal_text(doc, content)

    # Remove Synopsis section completely

    # Description always
    add_section("Description", data.get("description", ""))

    if "see_also" in data:
        add_section("See Also", data["see_also"])

    if "solution" in data:
        add_section("Solution", data["solution"])

    if "risk_factor" in data:
        add_section("Risk Factor", data["risk_factor"])

    if "cvss_v3.0_base_score" in data:
        add_section("CVSS v3.0 Base Score", data["cvss_v3.0_base_score"])

    if "cvss_v2.0_base_score" in data:
        add_section("CVSS v2.0 Base Score", data["cvss_v2.0_base_score"])

    if "plugin_information" in data:
        add_section("Plugin Information", data["plugin_information"])

    if "plugin_output" in data:
        # POC heading and image with minimal gap
        p_poc_title = doc.add_paragraph("Proof of Concept / Code Output")
        p_poc_title.style.font.name = 'Arial'
        p_poc_title.style.font.size = Pt(10)
        p_poc_title.space_after = Pt(2)  # minimal gap after heading
        p_poc_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        if image_path and os.path.exists(image_path):
            doc.add_picture(image_path, width=Inches(6))

# --- Add page footer with host left and page number right ---
def add_footer(doc, host_text):
    # Get section (assumes single section)
    section = doc.sections[0]

    # Create footer if not exists
    footer = section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()

    # Clear existing footer content
    footer_para.clear()

    # Add top border line (light gray)
    p = footer_para._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '4')
    top.set(qn('w:space'), '1')
    top.set(qn('w:color'), 'D9D9D9')  # light gray
    pBdr.append(top)
    pPr.append(pBdr)

    # Add two runs: host on left, page number on right
    run_left = footer_para.add_run(host_text)
    run_left.font.name = 'Arial'
    run_left.font.size = Pt(9)
    run_left.font.color.rgb = RGBColor(128, 128, 128)  # gray
    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Add tab to right with page number field
    # Use tab stops to push page number to the right margin
    tab_stops = footer_para.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(16))  # approx right margin

    footer_para.clear()
    footer_para.paragraph_format.tab_stops.add_tab_stop(Cm(16))

    # Construct text with tab character for spacing
    footer_para.add_run(host_text)
    footer_para.add_run('\t')

    # Add page number field
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.text = " PAGE "

    fldChar_separate = OxmlElement('w:fldChar')
    fldChar_separate.set(qn('w:fldCharType'), 'separate')

    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')

    r = footer_para.add_run()
    r._r.append(fldChar_begin)
    r._r.append(instrText)
    r._r.append(fldChar_separate)
    r._r.append(fldChar_end)

    # Apply font to entire paragraph (optional)
    for run in footer_para.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)

# --- MAIN SCRIPT ---

driver = webdriver.Chrome(options=options)
wait = WebDriverWait(driver, 10)
driver.get(URL)

doc = Document()
# Set margins
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Set default style font to Arial
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

try:
    time.sleep(2)

    if 'Login' in driver.title:
        username_field = driver.find_element(By.CLASS_NAME, 'login-username')
        password_field = driver.find_element(By.CLASS_NAME, 'login-password')
        remember_me = driver.find_element(By.CLASS_NAME, 'login-remember')
        sign_in_button = driver.find_element(By.TAG_NAME, 'button')

        username_field.clear()
        username_field.send_keys('admin')  # Change your username here
        password_field.send_keys('admin')  # Change your password here
        remember_me.click()
        sign_in_button.click()

    element = wait.until(EC.visibility_of_element_located((By.CLASS_NAME, 'select2-selection__rendered')))

    show_200_results_span = driver.find_element(By.CLASS_NAME, 'select2-selection__rendered')
    show_200_results_span.click()
    time.sleep(1)
    option_200 = driver.find_elements(By.CLASS_NAME, 'select2-results__option')[-1]
    option_200.click()
    time.sleep(1)

    list_of_vulnerabilities = driver.find_elements(By.CSS_SELECTOR, 'tr.vulnerability.add-plugin-id-tip')
    urls_to_visit = []

    for vulnerability in list_of_vulnerabilities:
        urls_to_visit.append(f"{URL}/{vulnerability.get_attribute('data-id')}")

    for index, url in enumerate(urls_to_visit):
        data_dictionary = {}

        driver.get(url)
        element = wait.until(EC.visibility_of_element_located((By.CLASS_NAME, 'add-plugin-id-tip')))

        section_data = driver.find_element(By.ID, 'content')
        section_html = section_data.get_attribute('innerHTML')

        soup = BeautifulSoup(section_html, 'html.parser')
        plugin_id = soup.find('tr', class_='noaction odd').get('data-plugin-id')
        vuln_severity = soup.find('span', class_='severity').text
        vuln_title = soup.find('h4', class_='add-plugin-id-tip').text

        data_dictionary['plugin_id'] = plugin_id
        data_dictionary['severity'] = vuln_severity
        data_dictionary['title'] = vuln_title

        vuln_data = soup.find_all('div', 'plugin-details-content')[0]

        # Normalize keys by replacing spaces with underscores and lowercase
        for header, content in zip(vuln_data.find_all('h5'), vuln_data.find_all('div', 'plugin-wrap')):
            key = header.text.lower().replace(" ", "_")
            data_dictionary[key] = content.text.strip()

        plugin_output_code = soup.find('pre', class_='monospace').text
        data_dictionary['plugin_output'] = plugin_output_code

        port_and_host = soup.find('tr', class_='noaction odd').find_all('td')
        port, host = port_and_host[0].text.strip(), port_and_host[1].text.strip()

        data_dictionary['port'] = port
        data_dictionary['host'] = host

        # Save image for code output
        image_filename = f"{plugin_id}.png"
        codeblock_image = text_to_image(plugin_output_code)
        codeblock_image.save(image_filename)

        # Append vulnerability to Word doc
        add_vulnerability_to_word_nessus_style(doc, data_dictionary, image_path=image_filename)

        # Add footer on every page with host and page number
        add_footer(doc, host)

        # Clean up image file after use
        os.remove(image_filename)

        print(f"Processed vulnerability {index+1}/{len(urls_to_visit)}", end='\r', flush=True)

    # Save final Word document
    doc.save("Nessus_Report.docx")
    print("\nReport saved as Nessus_Report.docx")

except Exception as e:
    print(f"An error occurred: {e}")

finally:
    driver.quit()
